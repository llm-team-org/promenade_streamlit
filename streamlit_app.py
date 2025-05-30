import streamlit as st
from prom_functions import (
    generate_company_information,
    generate_corp_code,
    short_list,
    sec_search,
    sec_get_report,
    dart_search,
    dart_get_report,
    get_dart_company_information
)
import asyncio
import os
import nest_asyncio
import tempfile
import traceback
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# Apply nest_asyncio to handle asyncio in Streamlit (as in original)
# This should be called once at the beginning.
nest_asyncio.apply()

# Configure Streamlit page
st.set_page_config(
    page_title="IM Draft Generator",
    page_icon="📊",
    layout="wide"
)

if 'report_list' not in st.session_state:
    st.session_state.report_list = []

if 'report_to_display' not in st.session_state:
    st.session_state.report_to_display = None

st.title("📊IM Draft Generator")
st.markdown("Generate comprehensive investment reports for companies using SEC or DART filings")

# Sidebar for inputs
st.sidebar.header("Configuration")

# Language selection
filings = st.sidebar.selectbox(
    "Select filings:",
    ["Global SEC filings", "Korean Dart fillings"],
    index=0
)

if filings == "Global SEC filings":
    language = "english"
else:
    language = "korean"

# Company URL input
company_url = st.sidebar.text_input(
    "Enter Company URL:",
    placeholder="https://example.com"
)

# Generate report button
generate_button = st.sidebar.button("🚀 Generate Report", type="primary")

if st.session_state.report_to_display:
    report_to_display_state = None
    st.sidebar.button(
        "View Instructions",
        on_click=lambda: setattr(st.session_state, 'report_to_display', None)
    )

# Display button for each report in st.session_state.report_list which when clicked will set the report_to_display variable to the report
if not st.session_state.report_list:
    st.sidebar.write("No reports found.")
else:
    st.sidebar.header("Report List")


def set_report_to_display(report):
    st.session_state.report_to_display = report


# Fixed section with unique keys
for i, report_data in enumerate(st.session_state.report_list):
    col1, col2 = st.sidebar.columns(2)
    col1.button(
        f"{report_data['url']}-{report_data['language']}",
        on_click=set_report_to_display,
        args=(report_data,),
        key=f"view_report_{i}"  # Unique key for each view button
    )
    col2.button(
        "❌",
        on_click=lambda report=report_data: st.session_state.report_list.remove(report),
        key=f"delete_report_{i}"  # Unique key for each delete button
    )


def add_dart_company_table(doc, corp_code_data):
    """Add DART company information table to the document"""
    # Add a heading for company information
    doc.add_heading('Company Information (DART)', level=1)

    # Create a table with 2 columns (Field, Value)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Add header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Value'

    # Make header bold - simplified approach
    for cell in hdr_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Only try to make bold if there are runs
        if cell.paragraphs[0].runs:
            for run in cell.paragraphs[0].runs:
                run.bold = True

    # Field mappings for better display names
    field_mappings = {
        'status': 'Status',
        'message': 'Message',
        'corp_code': 'Corporation Code',
        'corp_name': 'Corporation Name',
        'corp_name_eng': 'Corporation Name (English)',
        'stock_name': 'Stock Name',
        'stock_code': 'Stock Code',
        'ceo_nm': 'CEO Name',
        'corp_cls': 'Corporation Class',
        'jurir_no': 'Juridical Number',
        'bizr_no': 'Business Registration Number',
        'adres': 'Address',
        'hm_url': 'Homepage URL',
        'ir_url': 'IR URL',
        'phn_no': 'Phone Number',
        'fax_no': 'Fax Number',
        'induty_code': 'Industry Code',
        'est_dt': 'Establishment Date',
        'acc_mt': 'Account Month'
    }

    # Debug: Print what corp_code_data actually contains
    print(f"DEBUG: corp_code_data = {corp_code_data}")
    print(f"DEBUG: type(corp_code_data) = {type(corp_code_data)}")

    # Check if corp_code is 'N/A' OR empty dict - if so, set all values to 'N/A'
    if corp_code_data == 'N/A' or corp_code_data == {} or not corp_code_data:
        print("Adding N/A rows to table...")  # Debug print
        # Add all field mappings with 'N/A' values
        for key, field_name in field_mappings.items():
            row_cells = table.add_row().cells
            row_cells[0].text = field_name
            row_cells[1].text = 'N/A'

            # Center align the first column
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        print(f"Total rows in table: {len(table.rows)}")  # Debug print
    else:
        print("Adding regular data rows to table...")  # Debug print
        print(
            f"DEBUG: corp_code_data keys = {list(corp_code_data.keys()) if hasattr(corp_code_data, 'keys') else 'No keys method'}")

        # Add data rows normally
        rows_added = 0
        for key, value in corp_code_data.items():
            print(f"DEBUG: Checking key '{key}' - in field_mappings: {key in field_mappings}")
            if key in field_mappings:  # Only add mapped fields
                row_cells = table.add_row().cells
                row_cells[0].text = field_mappings[key]
                row_cells[1].text = str(value) if value else 'N/A'

                # Center align the first column
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                rows_added += 1
                print(f"DEBUG: Added row for key '{key}' with value '{value}'")

        print(f"DEBUG: Total rows added: {rows_added}")
        print(f"Total rows in table: {len(table.rows)}")  # Debug print

    # Add some spacing after the table
    doc.add_paragraph()
    doc.add_page_break()


def markdown_to_docx(markdown_text, company_name, language="english", corp_code_data=None):
    """Convert markdown text to a Word document"""
    doc = Document()

    # Add title
    title = doc.add_heading(f'Information Memorandum - {company_name}', 0)

    # Add DART company table if it's Korean language
    # Modified condition to handle both dict and 'N/A' string cases
    if language.lower() == "korean" and corp_code_data is not None:
        # Check if corp_code_data is a dict without error, or if it's 'N/A'
        if (isinstance(corp_code_data, dict) and "error" not in corp_code_data) or corp_code_data == 'N/A' or corp_code_data == {}:
            add_dart_company_table(doc, corp_code_data)

    # Split markdown into lines and process
    lines = markdown_text.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Handle headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        # Handle bullet points
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        # Handle numbered lists
        elif line.startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
            doc.add_paragraph(line[3:], style='List Number')
        # Handle bold text (basic implementation)
        elif '**' in line:
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    p.add_run(part)
                else:
                    p.add_run(part).bold = True
        # Regular paragraph
        else:
            doc.add_paragraph(line)

    return doc


def display_report(report_data):
    company_data = report_data.get('company_data', {})
    if not company_data or "error" in company_data:
        st.error(f"❌ Failed to extract company information: {company_data.get('error', 'Unknown error')}")
        if "raw_content" in company_data: st.expander("Raw LLM Output").write(company_data["raw_content"])
        return

    st.success("✅ Company information extracted successfully!")

    # Display company data
    st.subheader("📋 Company Information")
    with st.expander("View Company Details", expanded=True):
        st.json(company_data)

    # Extract key information
    full_name = company_data.get('company_name', 'N/A')
    first_name = company_data.get('company_first_name', 'N/A')
    selected_language = report_data.get('language', '')

    # Conditional display based on language
    if selected_language.lower() == "english":
        ticker = company_data.get('ticker', 'N/A')
        # Display basic info for English/SEC
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Company Name", full_name)
        with col2:
            st.metric("First Name", first_name)
        with col3:
            st.metric("Ticker", ticker)
    else:
        # For Korean/DART, show corp code instead of ticker
        corp_code_data = report_data.get('corp_code_data', {})
        corp_code = corp_code_data.get('corp_code', 'N/A') if corp_code_data else 'N/A'

        # Display basic info for Korean/DART
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Company Name", full_name)
        with col2:
            st.metric("First Name", first_name)
        with col3:
            st.metric("Corp Code", corp_code)

    if full_name == 'N/A':
        st.error("❌ Company name could not be determined. Cannot proceed.")
        return

    st.markdown("---")
    report = report_data.get('report', '')
    images = report_data.get('images', [])

    # Process based on language
    if selected_language.lower() == "english":
        st.subheader("🇺🇸 SEC Filing Analysis")

        filings_data = report_data.get('filings_data', {})

        if not filings_data or not filings_data.get('filings'):
            st.warning("⚠️ No SEC filings found or error in fetching.")
        else:
            st.success(f"✅ Found {len(filings_data.get('filings', []))} SEC filings.")
            with st.expander("View SEC Filings", expanded=False):
                st.json(filings_data)

            urls = [filing['filingUrl'] for filing in filings_data['filings'] if 'filingUrl' in filing]
            if not urls:
                st.warning("⚠️ No URLs found in SEC filings to generate report from.")
            else:
                st.success("✅ SEC report generated successfully!")

    elif selected_language.lower() == "korean":
        st.subheader("🇰🇷 DART Filing Analysis")

        corp_short_list_data = report_data.get('corp_short_list_data', {})
        report_source = report_data.get('report_source', 'web')
        web_search_reason = report_data.get('web_search_reason', '')

        if isinstance(corp_short_list_data, str) and "not in the dart list" in corp_short_list_data.lower():
            st.info("ℹ️ Company not in DART list. Report generated using web search.")
        elif web_search_reason == "not in short dart list":
            st.info("ℹ️ Company in DART list but not found in short DART list. Report generated using web search.")
        elif web_search_reason == "error in dart lookup":
            st.info("ℹ️ Error in DART lookup. Report generated using web search.")
        elif web_search_reason == "corp code generation failed":
            st.info("ℹ️ DART corporation code generation failed. Report generated using web search.")
        elif not corp_short_list_data:
            st.info("ℹ️ Company not found in DART list. Report generated using web search.")
        else:
            st.success("✅ Company found in DART short list.")
            with st.expander("View Short List", expanded=False):
                st.write(corp_short_list_data)

            corp_code_data = report_data.get('corp_code_data', {})
            if corp_code_data and "error" not in corp_code_data and corp_code_data.get('corp_code') != 'N/A':
                st.success("✅ DART Corporation code generated.")
                with st.expander("View Corporation Code", expanded=False):
                    st.json(corp_code_data)

        if report_source == 'web':
            st.info("ℹ️ Report generated using web search.")
        else:
            st.success("✅ Success! Report generated using DART filings!")

    # Commented out streaming logs display
    # if report_data.get('logs'):
    #     with st.expander("📊 Research Logs"):
    #         st.write(report_data['logs'])

    # Display report with download buttons
    if report:
        st.subheader(f"📈 {selected_language.capitalize()} Investment Report")

        # Add download buttons before the report display
        col1, col2, col3 = st.columns([1, 1, 3])

        # Create filename based on company name and language
        company_name_clean = full_name.replace(' ', '_').replace('/', '_').replace('\\', '_')

        with col1:
            filename_md = f"{company_name_clean}_{selected_language}_report.md"
            st.download_button(
                label="📥 Download MD",
                data=report,
                file_name=filename_md,
                mime="text/markdown",
                type="primary"
            )

        with col2:
            # Create Word document - pass corp_code_data for DART reports
            corp_code_data = report_data.get('corp_code_data', {}) if selected_language.lower() == "korean" else None
            doc = markdown_to_docx(report, full_name, selected_language, corp_code_data)

            # Save to bytes
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            filename_docx = f"{company_name_clean}_{selected_language}_report.docx"
            st.download_button(
                label="📄 Download DOCX",
                data=doc_buffer.getvalue(),
                file_name=filename_docx,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="secondary"
            )

        # Display the report
        st.markdown("---")
        st.markdown(report)

    elif not (
            "Error" in report if isinstance(report, str) else False):  # Only show if no error message already in report
        st.info("ℹ️ Report generation did not produce output, or path was skipped.")

    # Display images if any
    if images:
        st.subheader("🖼️ Report Images")
        for i, image_data in enumerate(images):
            # Assuming image_data could be URL, path, or bytes. St.image handles these.
            st.image(image_data, caption=f"Report Image {i + 1}")


# Main async function to handle the report generation flow
async def generate_report_flow(company_url_input, selected_language):
    report_data = {'url': company_url_input, 'language': selected_language}

    try:
        # Step 1: Generate company information
        with st.spinner("🔍 Analyzing company information..."):
            company_data = await generate_company_information(company_url_input, selected_language)
            report = company_data
            report_data['company_data'] = report

        if not company_data or "error" in company_data:
            st.error(f"❌ Failed to extract company information: {company_data.get('error', 'Unknown error')}")
            if "raw_content" in company_data:
                st.expander("Raw LLM Output").write(company_data["raw_content"])
            return

        st.success("✅ Company information extracted successfully!")

        # Display company data
        st.subheader("📋 Company Information")
        with st.expander("View Company Details", expanded=True):
            st.json(company_data)

        # Extract key information
        full_name = company_data.get('company_name', 'N/A')
        first_name = company_data.get('company_first_name', 'N/A')

        if full_name == 'N/A':
            st.error("❌ Company name could not be determined. Cannot proceed.")
            return

        # Conditional display based on language during generation
        if selected_language.lower() == "english":
            ticker = company_data.get('ticker', 'N/A')
            # Display basic info for English/SEC
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Company Name", full_name)
            with col2:
                st.metric("First Name", first_name)
            with col3:
                st.metric("Ticker", ticker)
        # For Korean/DART, we'll only show metrics after corp code is generated

        # Query template
        query_template = f"""As an investment associate, draft an information memorandum for company: {full_name}
        Information of Company: {company_data}

        -ADD These in table of contents:

        These are the Headings you need to use for IM and then generate sub headings for each heading
        1.Executive Summary
        2.Investment Highlights
        3.Company Overview
            Introduction to {full_name}
            History, Mission, and Core Values
            Global Presence and Operations
        4.Business Model, Strategy, and Product
        5.Business Segments Deep Dive
        6.Industry Overview and Competitive Positioning
        7.Financial Performance Analysis 
        8.Management and Corporate Governance
        9.Strategic Initiatives and Future Growth Drivers 
        10.Risk Factors 
        11.Investment Considerations
        12.Conclusion
        13.References
        
        -Add Tables: Display structured data like numbers, dates, comparisons, or lists in a table with headers, then summarize its main takeaways. For other content, use bullet points or numbered lists.

        -(Please exclude SWOT analysis)
        """
        st.markdown("---")
        report = ""
        images = []

        # Process based on language
        if selected_language.lower() == "english":
            st.subheader("🇺🇸 SEC Filing Analysis")

            try:
                with st.spinner("📄 Searching SEC filings..."):
                    filings_data = await sec_search(full_name, ticker)
                    report_data['filings_data'] = filings_data

                if not filings_data or not filings_data.get('filings'):
                    st.info("⚠️ No SEC filings found or error in fetching.")
                    urls = []
                else:
                    st.success(f"✅ Found {len(filings_data.get('filings', []))} SEC filings.")
                    with st.expander("View SEC Filings", expanded=False):
                        st.json(filings_data)

                    urls = [filing['filingUrl'] for filing in filings_data['filings'] if 'filingUrl' in filing]
                    if not urls:
                        st.warning("⚠️ No URLs found in SEC filings to generate report from.")

                # Add error handling for SEC report generation
                try:
                    with st.spinner("📊 Generating comprehensive IM report..."):
                        # Modified to not pass streaming containers - remove them entirely
                        report, images, logs = await sec_get_report(
                            query=query_template,
                            report_type="research_report",
                            sources=urls
                        )
                    report_data['report'] = report
                    report_data['images'] = images
                    # report_data['logs'] = logs  # Commented out logs
                    st.success("✅ IM report generated successfully!")
                except Exception as sec_error:
                    st.error(f"❌ Error generating report: {str(sec_error)}")
                    # Log the full error for debugging
                    st.expander("Error Details").write(f"Full error: {traceback.format_exc()}")
                    # Set fallback values
                    report_data['report'] = f"Error generating report: {str(sec_error)}"
                    report_data['images'] = []
                    # report_data['logs'] = f"Error occurred during report generation: {str(sec_error)}"

            except Exception as filing_error:
                st.error(f"❌ Error in SEC filing process: {str(filing_error)}")
                st.expander("Error Details").write(f"Full error: {traceback.format_exc()}")

        elif selected_language.lower() == "korean":
            st.subheader("🇰🇷 DART Filing Analysis")

            try:
                with st.spinner("📝 Generating company short list for DART..."):
                    company_first_name_for_dart = first_name if first_name != 'N/A' else full_name.split(" ")[0]
                    corp_short_list_data = await get_dart_company_information(full_name, company_first_name_for_dart)
                    # if corp_short_list_data == None:
                    #     corp_short_list_data = await short_list(full_name, company_first_name_for_dart)
                    report_data['corp_short_list_data'] = corp_short_list_data

                use_web_search = False
                web_search_reason = ""

                # Check if company is not in corp list at all
                if isinstance(corp_short_list_data, str) and "not in the dart list" in corp_short_list_data.lower():
                    st.info("ℹ️ Company not in DART list. Using web search instead.")
                    use_web_search = True
                    web_search_reason = "not in dart list"
                elif isinstance(corp_short_list_data, str) and "Error" in corp_short_list_data:
                    st.info(f"ℹ️ Error in DART lookup: {corp_short_list_data}. Using web search instead.")
                    use_web_search = True
                    web_search_reason = "error in dart lookup"
                elif not corp_short_list_data or (isinstance(corp_short_list_data, dict) and not corp_short_list_data):
                    st.info("ℹ️ Company in DART list but not found in short DART list. Using web search instead.")
                    use_web_search = True
                    web_search_reason = "not in short dart list"
                else:
                    st.success("✅ Company found in DART short list.")
                    with st.expander("View Short List", expanded=False):
                        st.write(corp_short_list_data)

                    with st.spinner("🔢 Generating DART corporation code..."):
                        corp_code_data = await generate_corp_code(full_name, corp_short_list_data, company_url_input)
                        st.write(corp_code_data)
                        if corp_code_data !='N/A':
                            corp_code_data=corp_short_list_data[int(corp_code_data)]
                            report_data['corp_code_data'] = corp_code_data

                            with st.expander("View Company Information", expanded=False):
                                st.write(corp_code_data)

                            with st.expander("View Corp Code", expanded=False):
                                st.write(corp_code_data['corp_code'])

                    if not corp_code_data or "error" in corp_code_data or corp_code_data == 'N/A':
                        st.info("ℹ️ Could not found company data in Dart Using web search instead.")
                        if "raw_content" in corp_code_data:
                            st.expander("Raw LLM Output").write(corp_code_data["raw_content"])
                        use_web_search = True
                        web_search_reason = "corp code generation failed"
                    else:
                        st.success("✅ DART Corporation code generated.")
                        # with st.expander("View Corporation Code", expanded=False):
                        #     st.write(corp_code_data)
                        corp_code_value = corp_code_data['corp_code']

                        # Show the complete metrics including corp code only once here
                        st.markdown("### 📊 Company Metrics")
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            st.metric("Company Name", full_name)
                        with col2:
                            st.metric("First Name", first_name)
                        with col3:
                            st.metric("Corp Code", corp_code_value)

                if use_web_search:
                    report_source = 'web'
                    report_data['report_source'] = report_source
                    report_data['web_search_reason'] = web_search_reason

                    try:
                        with st.spinner("📊 Generating IM report using web search..."):
                            report, images, logs = await dart_get_report(
                                query=query_template,
                                report_source=report_source,
                                path=None,
                                corp_data=corp_code_data
                            )
                        report_data['report'] = report
                        report_data['images'] = images
                        # report_data['logs'] = logs  # Commented out logs
                        st.success("✅ Report generated using web search!")
                    except Exception as dart_error:
                        st.error(f"❌ Error generating DART report (web search): {str(dart_error)}")
                        st.expander("Error Details").write(f"Full error: {traceback.format_exc()}")
                        report_data['report'] = f"Error generating report: {str(dart_error)}"
                        report_data['images'] = []
                        # report_data['logs'] = f"Error occurred during report generation: {str(dart_error)}"
                else:
                    st.info("✅ Company found in DART. Proceeding with DART filing download and report generation.")

                    with tempfile.TemporaryDirectory() as temp_dir:
                        try:
                            with st.spinner("📄 Searching DART filings and downloading documents..."):
                                doc_path = await dart_search(corp_code_value, temp_dir)

                            if not doc_path:
                                st.info("❌ Company data is not available in Dart documents. Using web sources instead.")
                                report_source = 'web'
                            else:
                                report_source = 'hybrid'
                                display_doc_path = os.path.relpath(doc_path, temp_dir)
                                st.success(f"✅ DART documents processed. Path: {display_doc_path}")
                                with st.expander("View Document Path", expanded=False):
                                    st.write(display_doc_path)

                            report_data['report_source'] = report_source

                            try:
                                with st.spinner("📊 Generating comprehensive IM report..."):
                                    report, images, logs = await dart_get_report(
                                        query=query_template,
                                        report_source=report_source,
                                        path=doc_path if report_source == 'hybrid' else None,
                                        corp_data=corp_code_data
                                    )
                                report_data['report'] = report
                                report_data['images'] = images
                                # report_data['logs'] = logs  # Commented out logs
                                st.success("✅ Success! Report generated using DART filings!")
                            except Exception as dart_filing_error:
                                st.error(f"❌ Error generating report: {str(dart_filing_error)}")
                                st.expander("Error Details").write(f"Full error: {traceback.format_exc()}")
                                report_data['report'] = f"Error generating report: {str(dart_filing_error)}"
                                report_data['images'] = []
                                # report_data['logs'] = f"Error occurred during report generation: {str(dart_filing_error)}"

                        except Exception as dart_search_error:
                            st.error(f"❌ Error in DART search process: {str(dart_search_error)}")
                            st.expander("Error Details").write(f"Full error: {traceback.format_exc()}")

            except Exception as dart_general_error:
                st.error(f"❌ Error in DART filing process: {str(dart_general_error)}")
                st.expander("Error Details").write(f"Full error: {traceback.format_exc()}")

        # Always add the report to session state and display, even if there were errors
        st.session_state.report_to_display = report_data
        st.session_state.report_list.append(report_data)

        # Display images if any
        if images:
            st.subheader("🖼️ Report Images")
            for i, image_data in enumerate(images):
                st.image(image_data, caption=f"Report Image {i + 1}")

    except Exception as general_error:
        st.error(f"❌ Unexpected error in report generation flow: {str(general_error)}")
        st.expander("Error Details").write(f"Full error: {traceback.format_exc()}")


# Main content area
if generate_button and company_url:
    try:
        if (company_url, language) in [(company_report["url"], company_report["language"]) for company_report in
                                       st.session_state.report_list]:
            st.info("⚠️ Report already generated for this company and language.")
        else:
            asyncio.run(generate_report_flow(company_url, language))
    except Exception as e:
        st.error(f"❌ An unexpected error occurred in the main application flow: {str(e)}")
        st.expander("Full Error Details").write(f"Complete traceback:\n{traceback.format_exc()}")

elif generate_button and not company_url:
    st.warning("⚠️ Please enter a company URL to generate the report.")

elif st.session_state.report_to_display:
    display_report(st.session_state.report_to_display)

else:
    # Welcome message (no changes here)
    st.markdown("""
    ## Welcome to the Information Memorendom Report Generator! 👋

    This application helps you generate comprehensive investment reports for companies using:

    - **SEC Filings** (for English/US companies)
    - **DART Filings** (for Korean companies)

    ### How to use:
    1. Select your preferred language (English or Korean)
    2. Enter the company's website URL
    3. Click "Generate Report" to start the analysis

    ### Features:
    - 🔍 Automatic company information extraction
    - 📄 Filing search and analysis
    - 📊 Comprehensive investment memorandum generation
    - 🖼️ Visual report elements (if generated)
    - 📥 Download reports as markdown files
    - 📄 Download reports as Word documents

    **Get started by entering a company URL in the sidebar!**
    """)

# Footer
st.markdown("---")
st.markdown(
    """
    <div style='text-align: center'>
        <p>IM Report Generator</p>
    </div>
    """,
    unsafe_allow_html=True)

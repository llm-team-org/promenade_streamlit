import streamlit as st
from prom_functions import (
    generate_company_information,
    generate_corp_code,
    short_list, # Note: short_list is imported in the new code but get_dart_company_information is used first for DART.
    sec_search,
    sec_get_report,
    dart_search,
    dart_get_report,
    get_dart_company_information # Ensure this function is defined in your prom_functions.py
)
import io
import asyncio
import os
import nest_asyncio
import tempfile
from google.genai import Client # Assuming this is the correct client for chat_object
from load_files import process_files_and_get_chat_object
from dotenv import load_dotenv
load_dotenv()
from sec_tool import sec_tool_function
from web_search import web_search_tool
from combined_tool import get_answer_to_query
import traceback
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
# Apply nest_asyncio to handle asyncio in Streamlit
nest_asyncio.apply()

# --- Page Constants ---
PAGE_REPORT_GENERATOR = "Report Generator"
PAGE_COMBINED_CHAT = "Chat With Tools"

# --- Page Configuration and Session State Initialization ---

def setup_page_config():
    """Sets up the Streamlit page configuration."""
    st.set_page_config(
        page_title="IM Draft Generator",
        page_icon="📊",
        layout="wide"
    )

def init_session_state():
    """Initializes session state variables if they don't exist."""
    if 'report_list' not in st.session_state:
        st.session_state.report_list = []
    if 'report_to_display' not in st.session_state:
        st.session_state.report_to_display = None
    if 'current_page' not in st.session_state:
        st.session_state.current_page = PAGE_REPORT_GENERATOR
    if 'uploaded_files' not in st.session_state: # Renamed from uploaded_pdfs
        st.session_state.uploaded_files = [{"name": "Web Search Tool", "path": "", "id": "Web Search Tool", 'tools_list': [web_search_tool]}, {"name":"SEC Filings Search Tool", "path":"", "id":"SEC Filings Search Tool", "tools_list":[sec_tool_function]}]
    if 'selected_file_for_chat' not in st.session_state: # Renamed from selected_pdf_for_chat
        st.session_state.selected_file_for_chat = None
    if 'last_filings_selection' not in st.session_state:
        st.session_state.last_filings_selection = "Global SEC filings"
    if 'last_company_url' not in st.session_state:
        st.session_state.last_company_url = ""
    if 'google_client' not in st.session_state:
        st.session_state.google_client = Client()
    # Removed 'chat_histories' as it was primarily used by the removed document_chat_page
    if 'sec_agent_query_answer' not in st.session_state: # Still used by combined_tools_chat_page
        st.session_state.sec_agent_query_answer = []
    # if 'available_tools' not in st.session_state:
        # st.session_state.available_tools = {"Web Search Tool":web_search_tool, "SEC Filings Search Tool": sec_tool_function}

# --- Helper Functions for UI and State Management ---

def write_multiline_text(text:str)->str:
    return "\n\n".join(text.splitlines())

def set_report_to_display(report):
    """Sets the report to be displayed in the main content area."""
    st.session_state.report_to_display = report

def remove_report_from_list(report_to_remove):
    """Removes a report from the report_list in session state."""
    st.session_state.report_list = [
        report for report in st.session_state.report_list if report != report_to_remove
    ]
    # If the removed report was currently displayed, clear the display
    if st.session_state.report_to_display == report_to_remove:
        st.session_state.report_to_display = None

def navigate_to(page_name):
    """Sets the current page in session state for navigation."""
    st.session_state.current_page = page_name
    st.session_state.report_to_display = None

# --- UI Rendering Functions ---

def markdown_to_docx(markdown_text, company_name, language="english", corp_code_data=None):
    """Convert markdown text to a Word document"""
    doc = Document()

    # Add title
    title = doc.add_heading(f'Information Memorandum - {company_name}', 0)

    # Add DART company table if it's Korean language
    # Modified condition to handle both dict and 'N/A' string cases
    #if language.lower() == "korean" and corp_code_data is not None:
        # # Check if corp_code_data is a dict without error, or if it's 'N/A'
        # if (isinstance(corp_code_data, dict) and "error" not in corp_code_data) or corp_code_data == 'N/A' or corp_code_data == {}:
        #     add_dart_company_table(doc, corp_code_data)

    # Split markdown into lines and process
    lines = markdown_text.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Handle headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        # Handle bullet points
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        # Handle numbered lists
        elif line.startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
            doc.add_paragraph(line[3:], style='List Number')
        # Handle bold text (basic implementation)
        elif '**' in line:
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    p.add_run(part)
                else:
                    p.add_run(part).bold = True
        # Regular paragraph
        else:
            doc.add_paragraph(line)

    return doc

def display_report(report_data):
    company_data = report_data.get('company_data', {})
    if not company_data or "error" in company_data:
        st.error(f"❌ Failed to extract company information: {company_data.get('error', 'Unknown error')}")
        if "raw_content" in company_data: st.expander("Raw LLM Output").write(write_multiline_text(company_data["raw_content"]))
        return

    st.success("✅ Company information extracted successfully!")

    # Display company data
    st.subheader("📋 Company Information")
    with st.expander("View Company Details", expanded=True):
        st.json(company_data)

    # Extract key information
    full_name = company_data.get('company_name', 'N/A')
    first_name = company_data.get('company_first_name', 'N/A')
    selected_language = report_data.get('language', '')

    # Conditional display based on language
    if selected_language.lower() == "english":
        ticker = company_data.get('ticker', 'N/A')
        # Display basic info for English/SEC
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Company Name", full_name)
        with col2:
            st.metric("First Name", first_name)
        with col3:
            st.metric("Ticker", ticker)
    else:
        # For Korean/DART, show corp code instead of ticker
        corp_code_data = report_data.get('corp_code_data', {})
        corp_code = corp_code_data.get('corp_code', 'N/A') if isinstance(corp_code_data, dict) else 'N/A'


        # Display basic info for Korean/DART
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Company Name", full_name)
        with col2:
            st.metric("First Name", first_name)
        with col3:
            st.metric("Corp Code", corp_code)

    if full_name == 'N/A':
        st.error("❌ Company name could not be determined. Cannot proceed.")
        return

    st.markdown("---")
    report = report_data.get('report', '')
    images = report_data.get('images', [])

    # Process based on language
    if selected_language.lower() == "english":
        st.subheader("🇺🇸 SEC Filing Analysis")

        filings_data = report_data.get('filings_data', {})

        if not filings_data or not filings_data.get('filings'):
            st.warning("⚠️ No SEC filings found or error in fetching.")
        else:
            st.success(f"✅ Found {len(filings_data.get('filings', []))} SEC filings.")
            with st.expander("View SEC Filings", expanded=False):
                st.json(filings_data)

            urls = [filing['filingUrl'] for filing in filings_data['filings'] if 'filingUrl' in filing]
            if not urls:
                st.warning("⚠️ No URLs found in SEC filings to generate report from.")
            else:
                st.success("✅ SEC report generated successfully!") # This message might be better placed after actual report generation step

    elif selected_language.lower() == "korean":
        st.subheader("🇰🇷 DART Filing Analysis")

        corp_short_list_data = report_data.get('corp_short_list_data', {})
        report_source = report_data.get('report_source', 'web')
        web_search_reason = report_data.get('web_search_reason', '')

        if isinstance(corp_short_list_data, str) and "not in the dart list" in corp_short_list_data.lower():
            st.info("ℹ️ Company not in DART list. Report generated using web search.")
        elif web_search_reason == "not in short dart list":
            st.info("ℹ️ Company in DART list but not found in short DART list. Report generated using web search.")
        elif web_search_reason == "error in dart lookup":
            st.info("ℹ️ Error in DART lookup. Report generated using web search.")
        elif web_search_reason == "corp code generation failed":
            st.info("ℹ️ DART corporation code generation failed. Report generated using web search.")
        elif not corp_short_list_data : # Handles empty dict or other falsy values for corp_short_list_data if it's not an error string
             if report_source == 'web': # Only show this if web search was indeed the fallback
                st.info("ℹ️ Company not found in DART list. Report generated using web search.")
        else: # corp_short_list_data has data
            st.success("✅ Company found in DART short list.")
            with st.expander("View Short List", expanded=False):
                st.write(corp_short_list_data)

            corp_code_data_from_report = report_data.get('corp_code_data', {}) # Renamed to avoid clash
            if isinstance(corp_code_data_from_report, dict) and "error" not in corp_code_data_from_report and corp_code_data_from_report.get('corp_code') != 'N/A':
                st.success("✅ DART Corporation code generated.")
                with st.expander("View Corporation Code Details", expanded=False): # Changed title for clarity
                    st.json(corp_code_data_from_report)

        if report_source == 'web':
            st.info("ℹ️ Report generated using web search.")
        else:
            st.success("✅ Success! Report generated using DART filings!")


    if report:
        st.subheader(f"📈 {selected_language.capitalize()} Investment Report")

        st.markdown("---")
        st.markdown(report)

    elif not ("Error" in report if isinstance(report, str) else False):
        st.info("ℹ️ Report generation did not produce output, or path was skipped.")

    if images:
        st.subheader("🖼️ Report Images")
        for i, image_data in enumerate(images):
            st.image(image_data, caption=f"Report Image {i + 1}")


async def generate_report_flow(company_url_input, selected_language):
    report_data = {'url': company_url_input, 'language': selected_language}

    try:
        with st.spinner("🔍 Analyzing company information..."):
            # Assuming generate_company_information is an async function from prom_functions
            company_data = await generate_company_information(company_url_input, selected_language)
            # The line `report = company_data` was present; unclear if intentional or a typo.
            # Storing company_data in report_data seems correct.
            report_data['company_data'] = company_data

        if not company_data or (isinstance(company_data, dict) and "error" in company_data):
            error_msg = company_data.get('error', 'Unknown error') if isinstance(company_data, dict) else "Invalid company data"
            st.error(f"❌ Failed to extract company information: {error_msg}")
            if isinstance(company_data, dict) and "raw_content" in company_data:
                st.expander("Raw LLM Output").write(company_data["raw_content"])
            return # Stop further processing

        st.success("✅ Company information extracted successfully!")

        st.subheader("📋 Company Information")
        with st.expander("View Company Details", expanded=True):
            st.json(company_data)

        full_name = company_data.get('company_name', 'N/A')
        first_name = company_data.get('company_first_name', 'N/A')

        if full_name == 'N/A':
            st.error("❌ Company name could not be determined. Cannot proceed.")
            return

        if selected_language.lower() == "english":
            ticker = company_data.get('ticker', 'N/A')
            col1, col2, col3 = st.columns(3)
            with col1: st.metric("Company Name", full_name)
            with col2: st.metric("First Name", first_name)
            with col3: st.metric("Ticker", ticker)
        # For Korean, metrics including corp_code are shown later after corp_code generation

        query_template = f"""As an investment associate, draft an information memorandum for company: {full_name}
        Information of Company: {company_data}

        -ADD These in table of contents:

        These are the Headings you need to use for IM and then generate sub headings for each heading
        1.Executive Summary
        2.Investment Highlights
        3.Company Overview
            Introduction to {full_name}
            History, Mission, and Core Values
            Global Presence and Operations
        4.Business Model, Strategy, and Product
        5.Business Segments Deep Dive
        6.Industry Overview and Competitive Positioning
        7.Financial Performance Analysis 
        8.Management and Corporate Governance
        9.Strategic Initiatives and Future Growth Drivers 
        10.Risk Factors 
        11.Investment Considerations
        12.Conclusion
        13.References
        
        -Add Tables: Display structured data like numbers, dates, comparisons, or lists in a table with headers, then summarize its main takeaways. For other content, use bullet points or numbered lists.

        -(Please exclude SWOT analysis)
        """
        st.markdown("---")
        report_content = "" # Renamed from 'report' to avoid conflict with company_data assignment earlier if it was a typo
        images = []
        corp_code_value = 'N/A' # Initialize for DART

        if selected_language.lower() == "english":
            st.subheader("🇺🇸 SEC Filing Analysis")
            try:
                with st.spinner("📄 Searching SEC filings..."):
                    ticker = company_data.get('ticker', 'N/A') # Ensure ticker is available
                    filings_data = await sec_search(full_name, ticker)
                    report_data['filings_data'] = filings_data

                if not filings_data or not filings_data.get('filings'):
                    st.info("⚠️ No SEC filings found or error in fetching.")
                    urls = []
                else:
                    st.success(f"✅ Found {len(filings_data.get('filings', []))} SEC filings.")
                    with st.expander("View SEC Filings", expanded=False):
                        st.json(filings_data)
                    urls = [filing['filingUrl'] for filing in filings_data['filings'] if 'filingUrl' in filing]
                    if not urls: st.warning("⚠️ No URLs found in SEC filings to generate report from.")

                try:
                    with st.spinner("📊 Generating comprehensive IM report..."):
                        report_content, images, _ = await sec_get_report( # Assuming logs are not needed here
                            query=query_template,
                            report_type="research_report",
                            sources=urls # Using all URLs as per new code
                        )
                    report_data['report'] = report_content
                    report_data['images'] = images
                    st.success("✅ IM report generated successfully!")
                except Exception as sec_error:
                    st.error(f"❌ Error generating report: {str(sec_error)}")
                    st.expander("Error Details").write(f"Full error: {write_multiline_text(traceback.format_exc())}")
                    return
                    report_data['report'] = f"Error generating report: {str(sec_error)}"
            except Exception as filing_error:
                st.error(f"❌ Error in SEC filing process: {str(filing_error)}")
                st.expander("Error Details").write(f"Full error: {write_multiline_text(traceback.format_exc())}")
                return
                report_data['report'] = f"Error in SEC filing process: {str(filing_error)}"


        elif selected_language.lower() == "korean":
            st.subheader("🇰🇷 DART Filing Analysis")
            corp_code_data_for_report = {} # Initialize
            try:
                with st.spinner("📝 Generating company short list for DART..."):
                    company_first_name_for_dart = first_name if first_name != 'N/A' else full_name.split(" ")[0]
                    # Using get_dart_company_information as per new script
                    corp_short_list_data = await get_dart_company_information(full_name, company_first_name_for_dart)
                    report_data['corp_short_list_data'] = corp_short_list_data

                use_web_search = False
                web_search_reason = ""

                if isinstance(corp_short_list_data, str) and "not in the dart list" in corp_short_list_data.lower():
                    st.info("ℹ️ Company not in DART list. Using web search instead.")
                    use_web_search = True
                    web_search_reason = "not in dart list"
                elif isinstance(corp_short_list_data, str) and "Error" in corp_short_list_data:
                    st.info(f"ℹ️ Error in DART lookup: {corp_short_list_data}. Using web search instead.")
                    use_web_search = True
                    web_search_reason = "error in dart lookup"
                elif not corp_short_list_data or (isinstance(corp_short_list_data, dict) and not corp_short_list_data):
                    st.info("ℹ️ Company in DART list but not found in short DART list. Using web search instead.")
                    use_web_search = True
                    web_search_reason = "not in short dart list"
                else: # Company found in DART short list (corp_short_list_data is likely a list of dicts)
                    st.success("✅ Company found in DART short list.")
                    with st.expander("View Short List", expanded=False): st.write(corp_short_list_data)

                    with st.spinner("🔢 Generating DART corporation code..."):
                        # generate_corp_code now takes company_url_input
                        selected_corp_index_str = await generate_corp_code(full_name, corp_short_list_data, company_url_input)
                        # st.write(selected_corp_index_str) # Original debug line

                        if selected_corp_index_str != 'N/A' and selected_corp_index_str is not None:
                            try:
                                selected_index = int(selected_corp_index_str)
                                if 0 <= selected_index < len(corp_short_list_data):
                                    corp_code_data_for_report = corp_short_list_data[selected_index]
                                    report_data['corp_code_data'] = corp_code_data_for_report
                                    corp_code_value = corp_code_data_for_report.get('corp_code', 'N/A')

                                    with st.expander("View Company Information (DART)", expanded=False):
                                        st.write(corp_code_data_for_report)
                                    with st.expander("View Corp Code (DART)", expanded=False):
                                        st.write(corp_code_value)
                                    st.success("✅ DART Corporation code processed.")
                                else:
                                    st.info("ℹ️ Invalid index for DART company. Using web search.")
                                    use_web_search = True
                                    web_search_reason = "corp code generation failed - invalid index"
                            except ValueError:
                                st.info("ℹ️ Corp code selection was not a valid number. Using web search.")
                                use_web_search = True
                                web_search_reason = "corp code generation failed - non-integer index"
                        else: # generate_corp_code returned 'N/A' or None
                             st.info("ℹ️ Could not determine company data in DART. Using web search instead.")
                             if isinstance(corp_code_data_for_report, dict) and "raw_content" in corp_code_data_for_report: # Check if corp_code_data_for_report got any raw_content
                                 st.expander("Raw LLM Output").write(corp_code_data_for_report["raw_content"])
                             use_web_search = True
                             web_search_reason = "corp code generation failed - N/A"

                # Display metrics for Korean company after attempting corp_code generation
                st.markdown("### 📊 Company Metrics (DART)")
                col1_k, col2_k, col3_k = st.columns(3)
                with col1_k: st.metric("Company Name", full_name)
                with col2_k: st.metric("First Name", first_name)
                with col3_k: st.metric("Corp Code", corp_code_value) # Shows N/A if not found

                if use_web_search:
                    report_source = 'web'
                    report_data['report_source'] = report_source
                    report_data['web_search_reason'] = web_search_reason
                    try:
                        with st.spinner("📊 Generating IM report using web search..."):
                            report_content, images, _ = await dart_get_report(
                                query=query_template, report_source=report_source, path=None
                            )
                        report_data['report'] = report_content
                        report_data['images'] = images
                        st.success("✅ Report generated using web search!")
                    except Exception as dart_web_error:
                        st.error(f"❌ Error generating DART report (web search): {str(dart_web_error)}")
                        st.expander("Error Details").write(f"Full error: {write_multiline_text(traceback.format_exc())}")
                        return
                        report_data['report'] = f"Error generating report (web): {str(dart_web_error)}"
                elif corp_code_value != 'N/A': # Proceed with DART documents only if corp_code was found
                    st.info("✅ Company found in DART. Proceeding with DART filing download and report generation.")
                    with tempfile.TemporaryDirectory() as temp_dir:
                        try:
                            with st.spinner("📄 Searching DART filings and downloading documents..."):
                                doc_path = await dart_search(corp_code_value, temp_dir)

                            if not doc_path:
                                st.info("❌ Company data is not available in DART documents. Using web sources instead.")
                                report_source = 'web' # Fallback to web
                                report_data['report_source'] = report_source
                                # Regenerate report with web source if docs not found
                                with st.spinner("📊 Generating IM report using web search (fallback)..."):
                                     report_content, images, _ = await dart_get_report(
                                        query=query_template, report_source='web', path=None)
                                     report_data['report'] = report_content
                                     report_data['images'] = images
                                     st.success("✅ Report generated using web search (fallback from no DART docs)!")

                            else: # Documents found
                                report_source = 'hybrid'
                                report_data['report_source'] = report_source
                                display_doc_path = os.path.relpath(doc_path, temp_dir)
                                st.success(f"✅ DART documents processed. Path: {display_doc_path}")
                                with st.expander("View Document Path", expanded=False): st.write(display_doc_path)

                                with st.spinner("📊 Generating comprehensive IM report from DART docs..."):
                                    report_content, images, _ = await dart_get_report(
                                        query=query_template, report_source=report_source, path=doc_path
                                    )
                                report_data['report'] = report_content
                                report_data['images'] = images
                                st.success("✅ Success! Report generated using DART filings!")

                        except Exception as dart_filing_error:
                            st.error(f"❌ Error generating report from DART filings: {str(dart_filing_error)}")
                            st.expander("Error Details").write(f"Full error: \n{write_multiline_text(traceback.format_exc())}")
                            return
                            report_data['report'] = f"Error generating report (DART filings): {str(dart_filing_error)}"
                else: # Not using web search but corp_code_value is N/A - this case should be handled by web_search_reason
                    st.warning("ℹ️ Could not proceed with DART document search as Corp Code was not identified.")
                    report_data['report'] = "Could not obtain DART Corp Code for document search."


            except Exception as dart_general_error:
                st.error(f"❌ Error in DART filing process: {str(dart_general_error)}")
                st.expander("Error Details").write(f"Full error: {write_multiline_text(traceback.format_exc())}")
                return
                report_data['report'] = f"Error in DART filing process: {str(dart_general_error)}"

        st.session_state.report_to_display = report_data
        if not any(existing_report['url'] == report_data['url'] and existing_report['language'] == report_data['language'] for existing_report in st.session_state.report_list):
            st.session_state.report_list.append(report_data)
        st.rerun()

        # This image display seems redundant if display_report is called immediately after.
        # However, if generate_report_flow is meant to update the main area directly:
        if images:
            st.subheader("🖼️ Report Images (from generation)")
            for i, image_data in enumerate(images):
                st.image(image_data, caption=f"Report Image {i + 1}")
        # Calling display_report here if this function is responsible for the final main page update
        # display_report(report_data) # Or rely on the main script logic to call display_report


    except Exception as general_error:
        st.error(f"❌ Unexpected error in report generation flow: {str(general_error)}")
        st.expander("Error Details").write(f"Full error: {write_multiline_text(traceback.format_exc())}")
        return
        # Ensure report_data has some error message if an overarching error occurs
        if 'report' not in report_data or not report_data['report'] :
             report_data['report'] = f"Unexpected error in report generation: {str(general_error)}"
        st.session_state.report_to_display = report_data # Display error info
        # Optionally add to list for review
        if not any(existing_report['url'] == report_data['url'] and existing_report['language'] == report_data['language'] for existing_report in st.session_state.report_list):
            st.session_state.report_list.append(report_data)

def display_report_details(report_data):
    """Displays the comprehensive report details in the main content area."""
    company_data = report_data.get('company_data', {})
    if not company_data or "error" in company_data:
        st.error(f"❌ Failed to extract company information: {company_data.get('error', 'Unknown error')}")
        if "raw_content" in company_data:
            st.expander("Raw LLM Output").write(company_data["raw_content"])
        return

    st.success("✅ Company information extracted successfully!")

    # Display company data
    st.subheader("📋 Company Information")
    with st.expander("View Company Details", expanded=True):
        st.json(company_data)

    full_name = company_data.get('company_name', 'N/A')
    first_name = company_data.get('company_first_name', 'N/A')
    ticker = company_data.get('ticker', 'N/A')

    if full_name == 'N/A':
        st.error("❌ Company name could not be determined. Cannot proceed.")
        return

    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Company Name", full_name)
    with col2:
        st.metric("First Name", first_name)
    with col3:
        st.metric("Ticker", ticker)

    st.markdown("---")

    selected_language = report_data.get('language', '')
    report = report_data.get('report', '')
    images = report_data.get('images', [])

    if selected_language.lower() == "english":
        st.subheader("🇺🇸 SEC Filing Analysis")
        filings_data = report_data.get('filings_data', {})

        if not filings_data or not filings_data.get('filings'):
            st.warning("⚠️ No SEC filings found or error in fetching.")
        else:
            st.success(f"✅ Found {len(filings_data.get('filings', []))} SEC filings.")
            with st.expander("View SEC Filings", expanded=False):
                st.json(filings_data)

    elif selected_language.lower() == "korean":
        st.subheader("🇰🇷 DART Filing Analysis")
        corp_short_list_data = report_data.get('corp_short_list_data', {})
        report_source = report_data.get('report_source', 'web')
        web_search_reason = report_data.get('web_search_reason', '')

        if isinstance(corp_short_list_data, str) and "not in the dart list" in corp_short_list_data.lower():
            st.info("ℹ️ Company not in DART list. Report generated using web search.")
        elif web_search_reason == "not in short dart list":
            st.info("ℹ️ Company in DART list but not found in short DART list. Report generated using web search.")
        elif web_search_reason == "error in dart lookup":
            st.info("ℹ️ Error in DART lookup. Report generated using web search.")
        elif web_search_reason == "corp code generation failed":
            st.info("ℹ️ DART corporation code generation failed. Report generated using web search.")
        elif not corp_short_list_data:
            st.info("ℹ️ Company not found in DART list. Report generated using web search.")
        else:
            st.success("✅ Company found in DART short list.")
            with st.expander("View Short List", expanded=False):
                st.write(corp_short_list_data)

            corp_code_data = report_data.get('corp_code_data', {})
            if corp_code_data and "error" not in corp_code_data and corp_code_data.get('corp_code') != 'N/A':
                st.success("✅ DART Corporation code generated.")
                with st.expander("View Corporation Code", expanded=False):
                    st.json(corp_code_data)

        if report_source == 'web':
            st.info("ℹ️ Report generated using web search.")
        else:
            st.success("✅ Success! Report generated using DART filings!")

    if report_data.get('logs'):
        with st.expander("📊 Research Logs"):
            st.write(report_data['logs'])

    if report:
        st.subheader(f"📈 {selected_language.capitalize()} Investment Report")
        company_name_clean = full_name.replace(' ', '_').replace('/', '_').replace('\\', '_')

        with st.expander("View Full Report", expanded=True):
            st.markdown(report)
    else:
        st.info("ℹ️ Report generation did not produce output, or path was skipped.")

    if images:
        st.subheader("🖼️ Report Images")
        for i, image_data in enumerate(images):
            st.image(image_data, caption=f"Report Image {i + 1}")

def display_welcome_message():
    """Displays the welcome message and instructions."""
    st.markdown("""
    ## Welcome to the Investment Report Generator! 👋

    This application helps you generate comprehensive investment reports for companies using:

    - **SEC Filings** (for English/US companies)
    - **DART Filings** (for Korean companies)

    ### How to use:
    1. Select your preferred **filings type** (Global SEC or Korean DART) above.
    2. Enter the **company's website URL** in the input field.
    3. Click "**🚀 Generate Report**" to start the analysis.

    ### Features:
    - 🔍 Automatic company information extraction
    - 📄 Filing search and analysis
    - 📊 Comprehensive investment memorandum generation
    - 🖼️ Visual report elements (if generated)
    - 📥 Download reports as markdown files

    **Get started by filling out the details above!**
    """)


# --- Helper Functions for Agent Logic ---

def process_uploaded_file(uploaded_file):
    """
    Processes the uploaded file (e.g., extracts text, creates embeddings)
    and initializes a chat object.
    Returns a dictionary containing file info and the chat object.
    """
    if uploaded_file:
        file_extension = os.path.splitext(uploaded_file.name)[1]
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            file_path = tmp_file.name
            tools_list = process_files_and_get_chat_object(file_path_list=[file_path], client=st.session_state.google_client)
        st.success(f"Document '{uploaded_file.name}' processed and ready for chat.")
        return {"name": uploaded_file.name, "path": file_path, "id": os.path.basename(file_path), 'tools_list': tools_list}
    return None

# --- UI Rendering Functions (Pages) ---
def render_sidebar_navigation():
    """Renders the navigation buttons in the sidebar."""
    st.sidebar.header("Navigation")
    if st.sidebar.button("📊 Report Generator", key="nav_report_gen"):
        navigate_to(PAGE_REPORT_GENERATOR)
    if st.sidebar.button("🛠️ Chat with Tools",key="nav_chat_tools"):
        navigate_to(PAGE_COMBINED_CHAT)

def render_report_generator_page():
    """Renders the main Report Generator page content."""
    st.title("📊IM Draft Generator")
    st.markdown("Generate comprehensive investment reports for companies using SEC or DART filings")

    st.markdown("---")

    # --- Section: Generate New Report ---
    st.header("✨ Generate New Report")
    with st.container(border=True):
        col1, col2 = st.columns([1, 2])

        with col1:
            filings_selection = st.selectbox(
                "Select filings:",
                ["Global SEC filings", "Korean Dart fillings"],
                index=["Global SEC filings", "Korean Dart fillings"].index(st.session_state.last_filings_selection)
            )
            language = "english" if filings_selection == "Global SEC filings" else "korean"
            st.session_state.last_filings_selection = filings_selection

        with col2:
            company_url = st.text_input(
                "Enter Company URL:",
                value=st.session_state.last_company_url,
                placeholder="https://example.com"
            )
            st.session_state.last_company_url = company_url

        generate_button = st.button("🚀 Generate Report", type="primary")

    if generate_button:
        if not company_url:
            st.warning("⚠️ Please enter a company URL to generate the report.")
        else:
            is_duplicate = any(
                r["url"] == company_url and r["language"] == language
                for r in st.session_state.report_list
            )
            if is_duplicate:
                st.info("⚠️ A report for this company and language has already been generated. Displaying the existing report.")
                existing_report = next(
                    (r for r in st.session_state.report_list if r["url"] == company_url and r["language"] == language),
                    None
                )
                if existing_report:
                    set_report_to_display(existing_report)
            else:
                try:
                    asyncio.run(generate_report_flow(company_url, language))
                except Exception as e:
                    st.error(f"❌ An unexpected error occurred during report generation: {str(e)}")
                    st.exception(e)

    st.markdown("---")

    # --- Section: Generated Reports ---
    st.header("📄 Generated Reports")
    if not st.session_state.report_list:
        st.info("No reports generated yet. Use the section above to create one!")
    else:
        for i, report_data in enumerate(st.session_state.report_list):
            col1,col2, col3,col4 = st.columns([3,1,1,1])
            with col1:
                company_full_name = report_data['company_data']['company_name'].replace(' ', '_').replace('/', '_').replace('\\', '_')
                st.button(
                    f"View {company_full_name}_{report_data['language']} Report",
                    on_click=set_report_to_display,
                    args=(report_data,),
                    key=f"view_report_{i}",
                    type="primary",
                    use_container_width=True
                )
            with col2:
                company_full_name = report_data['company_data']['company_name'].replace(' ', '_').replace('/', '_').replace('\\', '_')
                filename = f"{company_full_name}_{report_data['language']}_report.md"
                st.download_button(
                    label="📥 Download MD",
                    key=f"download_report_{i}",
                    data=report_data['report'],
                    file_name=filename,
                    mime="text/markdown",
                    use_container_width=True
                )
            with col3:
                company_full_name = report_data['company_data']['company_name'].replace(' ', '_').replace('/', '_').replace('\\', '_')
                report_text = report_data['report']
                selected_language = report_data['language']
                corp_code_data = report_data.get('corp_code_data', {}) if selected_language.lower() == "korean" else None
                doc = markdown_to_docx(report_text, company_full_name, selected_language, corp_code_data)

                # Save to bytes
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)

                filename_docx = f"{company_full_name}_{selected_language}_report.docx"
                st.download_button(
                    label="📄 Download DOCX",
                    data=doc_buffer.getvalue(),
                    file_name=filename_docx,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="secondary",
                    use_container_width=True
                )
            with col4:
                st.button(
                    "❌ Remove",
                    on_click=remove_report_from_list,
                    args=(report_data,),
                    key=f"delete_report_{i}",
                    use_container_width=True
                )
    st.markdown("---")

    if st.session_state.report_to_display:
        st.header("📊 Current Report Details")
        display_report(st.session_state.report_to_display)
        if st.button("Clear Report Display", help="Click to hide the currently displayed report details."):
            set_report_to_display(None)
    elif not generate_button and not st.session_state.report_to_display:
        display_welcome_message()

def combined_tools_chat_page():
    st.markdown("Upload a document and chat with its content using available tools.")

    st.subheader("Upload Document")
    uploaded_file = st.file_uploader(
        "Choose a file to add its specific chat tool",
        type=["pdf", "txt", "csv", "docx", "xlsx"],
        key="file_uploader_combined" # Changed key to avoid conflict if other uploaders exist
    )
    if uploaded_file:
        existing_file_names = [file['name'] for file in st.session_state.uploaded_files]
        if uploaded_file.name not in existing_file_names:
            with st.spinner(f"Processing {uploaded_file.name}..."):
                processed_info = process_uploaded_file(uploaded_file)
                if processed_info:
                    st.session_state.uploaded_files.append(processed_info)
                    # No need to initialize chat_histories here as it's removed
                else:
                    st.error("Failed to process document.")
        else:
            st.info(f"Document '{uploaded_file.name}' is already uploaded and processed.")
            # Logic to select this file for chat could be added if needed

    st.subheader("Chat with Tools")
    # Ensure default tools are always options and selected by default
    default_tool_names = ["Web Search Tool", "SEC Filings Search Tool"]
    available_tool_options = [file['name'] for file in st.session_state.uploaded_files]
    
    # Determine default selections: all default tools + any newly uploaded unique tools
    current_selection = [tool_name for tool_name in default_tool_names if tool_name in available_tool_options]
    # Add other uploaded files to selection if they are not already part of default
    for file_info in st.session_state.uploaded_files:
        if file_info['name'] not in current_selection:
            current_selection.append(file_info['name'])


    selected_tools_names = st.multiselect(
        "Select the tools you want to use for your query:",
        options=available_tool_options,
        default=current_selection # Select all available tools by default
        )

    # Display previous Q&A using st.session_state.sec_agent_query_answer
    for query_answer in st.session_state.sec_agent_query_answer:
        with st.container(border=True):
            with st.chat_message('user'):
                st.write(query_answer['query'])
            with st.chat_message('assistant'):
                st.write(query_answer['answer'])
    
    user_query = st.chat_input("Ask your query using the selected tools:")
    if user_query:
        if not selected_tools_names:
            st.warning("Please select at least one tool to query.")
        else:
            with st.container(border=True):
                with st.chat_message("user"):
                    st.write(user_query)
                with st.spinner("Getting answer using selected tools..."):
                    # Flatten the list of tool lists
                    tools_list_to_send = []
                    for file_info in st.session_state.uploaded_files:
                        if file_info['name'] in selected_tools_names:
                            tools_list_to_send.extend(file_info['tools_list'])
                    
                    answer = get_answer_to_query(user_query, tools_list_to_send) # Assuming get_answer_to_query can handle a flat list of tools
                if not answer:
                    answer = "Failed to get answer from the combined tools."
                new_query_answer = {'query':user_query, 'answer':answer}
                st.session_state.sec_agent_query_answer.append(new_query_answer)
                st.rerun()

# --- Core Report Generation Logic (Unchanged) ---

async def generate_report_flow_async(company_url_input, selected_language):
    """
    Handles the asynchronous flow of generating an investment report.
    Fetches company info, then proceeds with SEC or DART filings.
    """
    report_data = {'url': company_url_input, 'language': selected_language}
    st.session_state.report_to_display = None # Clear previous display during generation

    # Step 1: Generate company information
    with st.spinner("🔍 Analyzing company information..."):
        company_data = await generate_company_information(company_url_input, selected_language)
        report_data['company_data'] = company_data

    if not company_data or "error" in company_data:
        st.error(f"❌ Failed to extract company information: {company_data.get('error', 'Unknown error')}")
        if "raw_content" in company_data:
            st.expander("Raw LLM Output").write(company_data["raw_content"])
        return

    st.success("✅ Company information extracted successfully!")

    full_name = company_data.get('company_name', 'N/A')
    first_name = company_data.get('company_first_name', 'N/A')

    if full_name == 'N/A':
        st.error("❌ Company name could not be determined. Cannot proceed.")
        return

    english_query_template = f"""As an investment associate, draft an information memorandum for company: {full_name}
    Information of Company: {company_data}
    ADD These in table of contents:

    These are the Headings you need to use for IM and then generate sub headings for each heading
    1.Executive Summary
    2.Investment Highlights
    3.Company Overview
        Introduction to {full_name}
        History, Mission, and Core Values
        Global Presence and Operations
    4.Business Model, Strategy, and Product
    5.Business Segments Deep Dive
    6.Industry Overview and Competitive Positioning
    7.Financial Performance Analysis
    8.Management and Corporate Governance
    9.Strategic Initiatives and Future Growth Drivers
    10.Risk Factors
    11.Investment Considerations
    12.Conclusion
    13.References
    """

    korean_query_template = f"""투자 담당자로서 회사 {full_name}에 대한 정보 메모를 작성하십시오.
    회사 정보: {company_data}
    목차에 다음 내용을 추가하십시오.

    정보 메모에 사용해야 하는 제목은 다음과 같으며, 각 제목에 대한 하위 제목을 생성해야 합니다.
    1. 요약
    2. 투자 주요 내용
    3. 회사 개요
    {full_name} 소개
    연혁, 사명 및 핵심 가치
    글로벌 입지 및 운영
    4. 비즈니스 모델, 전략 및 제품
    5. 사업 부문 심층 분석
    6. 산업 개요 및 경쟁적 포지셔닝
    7. 재무 성과 분석
    8. 경영진 및 기업 지배구조
    9. 전략적 이니셔티브 및 미래 성장 동력
    10. 위험 요소
    11. 투자 고려 사항
    12. 결론
    13. 참고 문헌
    """
    report = ""
    images = []
    logs = ""

    with st.expander('📊 Research Logs', expanded=True):
        logs_container = st.empty()
        logs_container.info("Logs will appear here as the research progresses...")
    reports_container = st.empty()
    reports_container.info("The final report will be displayed here once generated.")

    if selected_language.lower() == "english":
        st.info("Searching SEC filings...")
        with st.spinner("📄 Searching SEC filings..."):
            filings_data = await sec_search(full_name)
            report_data['filings_data'] = filings_data

        if not filings_data or not filings_data.get('filings'):
            st.warning("⚠️ No SEC filings found or error in fetching.")
            urls = []
        else:
            st.success(f"✅ Found {len(filings_data.get('filings', []))} SEC filings.")
            urls = [filing['filingUrl'] for filing in filings_data['filings'] if 'filingUrl' in filing]
            if not urls:
                st.warning("⚠️ No URLs found in SEC filings to generate report from.")

        with st.spinner("📊 Generating comprehensive SEC report..."):
            report, images, logs = await sec_get_report(
                query=english_query_template,
                report_type="research_report",
                sources=urls[:2],
                logs_container=logs_container,
                report_container=reports_container
            )
        st.success("✅ SEC report generation complete!")

    elif selected_language.lower() == "korean":
        st.info("Initiating DART filing analysis...")
        use_web_search = False
        web_search_reason = ""
        corp_code_value = None

        with st.spinner("📝 Generating company short list for DART..."):
            company_first_name_for_dart = first_name if first_name != 'N/A' else full_name.split(" ")[0]
            corp_short_list_data = await short_list(full_name, company_first_name_for_dart)
            report_data['corp_short_list_data'] = corp_short_list_data

        if isinstance(corp_short_list_data, str) and "not in the dart list" in corp_short_list_data.lower():
            st.info("ℹ️ Company not in DART list. Using web search instead.")
            use_web_search = True
            web_search_reason = "not in dart list"
        elif isinstance(corp_short_list_data, str) and "Error" in corp_short_list_data:
            st.info(f"ℹ️ Error in DART lookup: {corp_short_list_data}. Using web search instead.")
            use_web_search = True
            web_search_reason = "error in dart lookup"
        elif not corp_short_list_data:
            st.info("ℹ️ Company in DART list but not found in short DART list. Using web search instead.")
            use_web_search = True
            web_search_reason = "not in short dart list"
        else:
            st.success("✅ Company found in DART short list.")
            with st.spinner("🔢 Generating DART corporation code..."):
                corp_code_data = await generate_corp_code(full_name, corp_short_list_data)
                report_data['corp_code_data'] = corp_code_data

            if not corp_code_data or "error" in corp_code_data or corp_code_data.get('corp_code') == 'N/A':
                st.info("ℹ️ Could not find company data in DART. Using web search instead.")
                if "raw_content" in corp_code_data:
                    st.expander("Raw LLM Output").write(corp_code_data["raw_content"])
                use_web_search = True
                web_search_reason = "corp code generation failed"
            else:
                st.success("✅ DART Corporation code generated.")
                corp_code_value = corp_code_data['corp_code']

        report_source = 'web' if use_web_search else 'hybrid'
        report_data['report_source'] = report_source
        report_data['web_search_reason'] = web_search_reason

        doc_path = None
        if not use_web_search:
            with tempfile.TemporaryDirectory() as temp_dir:
                with st.spinner("📄 Searching DART filings and downloading documents..."):
                    doc_path = await dart_search(corp_code_value, temp_dir)

                if not doc_path:
                    st.info("❌ Company data is not available in DART documents. Using web sources instead.")
                    report_source = 'web'
                    report_data['report_source'] = report_source
                else:
                    st.success(f"✅ DART documents processed.")

        with st.spinner(f"📊 Generating comprehensive DART report using {report_source} sources..."):
            report, images, logs = await dart_get_report(
                query=korean_query_template,
                report_source=report_source,
                path=doc_path,
                logs_container=logs_container,
                report_container=reports_container
            )
        st.success("✅ DART report generation complete!")

    report_data['report'] = report
    report_data['images'] = images
    report_data['logs'] = logs

    st.session_state.report_list.append(report_data)
    st.session_state.report_to_display = report_data
    st.rerun()


# --- Main Application Runner ---

def main():
    """Main function to run the Streamlit application."""
    setup_page_config()
    init_session_state()

    render_sidebar_navigation()

    if st.session_state.current_page == PAGE_REPORT_GENERATOR:
        render_report_generator_page()
    elif st.session_state.current_page == PAGE_COMBINED_CHAT:
        combined_tools_chat_page()

if __name__ == "__main__":
    main()